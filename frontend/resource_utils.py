"""Pure logic for the frontend, kept separate from app.py so it's importable
and testable without needing a running Streamlit script context.
"""

import html
import io
import json
import sys
from pathlib import Path

import markdown as md
from docx import Document
from htmldocx import HtmlToDocx
from xhtml2pdf import pisa

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from schemas.sections import LessonPlanSections  # noqa: E402

SCHOOL_CONFIG_PATH = Path(__file__).resolve().parent / "school_config.json"

# Lesson Plan's short identifying facts - rendered as one compact table
# instead of a heading + one-line paragraph each, which read as a wall of
# disconnected one-liners in the exported Word/PDF document.
LESSON_PLAN_INFO_FIELDS = ["grade", "subject", "strand", "unit", "topic"]


def _load_config() -> dict:
    if not SCHOOL_CONFIG_PATH.exists():
        return {}
    try:
        return json.loads(SCHOOL_CONFIG_PATH.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return {}


def _save_config(config: dict) -> None:
    SCHOOL_CONFIG_PATH.write_text(json.dumps(config), encoding="utf-8")


def load_school_name() -> str | None:
    return _load_config().get("school_name") or None


def save_school_name(name: str) -> None:
    config = _load_config()
    config["school_name"] = name
    _save_config(config)


def format_curriculum_source(curriculum_record: dict | None) -> str | None:
    """Builds a human-readable citation from a matched curriculum record's
    `source_refs`, e.g. for display under generated content.
    """
    if not curriculum_record:
        return None
    refs = curriculum_record.get("source_refs") or {}
    parts = []
    syllabus = refs.get("syllabus")
    if syllabus and syllabus.get("document"):
        pages = syllabus.get("content_standard_page")
        page_note = f" (p.{pages})" if pages else ""
        parts.append(f"{syllabus['document']}{page_note}")
    teacher_guide = refs.get("teacher_guide")
    if teacher_guide and teacher_guide.get("document"):
        parts.append(teacher_guide["document"])
    if not parts:
        return None
    return "; ".join(parts)


def render_package_markdown(package) -> str:
    """Renders a ResourcePackage into a single Markdown document for display
    and PDF/text export.
    """
    parts: list[str] = []

    if package.custom_topic_notice:
        parts.append(f"> **Note:** {package.custom_topic_notice}")

    if package.review.alignment_status == "needs_review":
        issues = "\n".join(f"- {issue}" for issue in package.review.issues)
        parts.append(f"> **This output was flagged for review:**\n{issues}")

    for title, section in (
        ("Lesson Plan", package.lesson_plan),
        ("Lesson Activities", package.activities),
        ("Teacher Notes", package.teacher_notes),
        ("Assessment", package.assessment),
    ):
        if section is None:
            continue
        parts.append(render_section_markdown(title, section))

    source = format_curriculum_source(package.match.curriculum_record)
    if source:
        parts.append(f"---\n\n**Source(s):**\n- {source}")

    return "\n\n".join(parts)


def render_section_markdown(title: str, section) -> str:
    """Renders one generated section (e.g. a parsed `LessonPlanSections`) as a
    `##`-headed Markdown block, used both per-wizard-step and inside
    `render_package_markdown`'s final combined document.
    """
    dumped = section.model_dump(exclude={"raw_markdown"})
    lines = [f"## {title}"]
    skip_fields: set[str] = set()

    if isinstance(section, LessonPlanSections):
        if dumped.get("lesson_title"):
            lines.append(f"**{dumped['lesson_title']}**")
        info_rows = [(name, dumped[name]) for name in LESSON_PLAN_INFO_FIELDS if dumped.get(name)]
        if info_rows:
            # Raw HTML, not a Markdown table - a Markdown table always
            # promotes its first row to <th>, which the app's CSS renders as
            # a solid black bar; this table has no header row to give it one.
            row_html = "".join(
                f"<tr><td><strong>{html.escape(name.title())}</strong></td><td>{html.escape(value)}</td></tr>"
                for name, value in info_rows
            )
            lines.append(f"<table>{row_html}</table>")
        skip_fields = {"lesson_title", *LESSON_PLAN_INFO_FIELDS}

    for field_name, value in dumped.items():
        if field_name in skip_fields or not value:
            continue
        heading = field_name.replace("_", " ").title()
        lines.append(f"### {heading}\n\n{value}")
    return "\n\n".join(lines)


def markdown_to_pdf_bytes(markdown_text: str, title: str, school_name: str | None = None) -> bytes:
    # Visually separate the Source(s) section, generated as a plain bold
    # paragraph by the markdown converter, with a rule so it reads as a
    # distinct footnote-style block rather than just another paragraph.
    text = markdown_text.replace("**Source(s):**", "---\n\n**Source(s):**")
    body_html = md.markdown(text, extensions=["tables", "nl2br"])

    header_band = ""
    if school_name:
        header_band = f"""
        <table class="header-band"><tr><td>
            <div class="school-name">{school_name}</div>
            <div class="app-name">PNG Classroom Resource Generator</div>
        </td></tr></table>
        """

    html = f"""
    <html>
    <head>
    <style>
        @page {{
            size: A4;
            margin: 2cm 1.8cm 2.2cm 1.8cm;
            @frame footer_frame {{
                -pdf-frame-content: footer_content;
                bottom: 0.8cm; margin-left: 1.8cm; margin-right: 1.8cm; height: 1cm;
            }}
        }}
        body {{ font-family: Helvetica, Arial, sans-serif; font-size: 10.5pt; color: #1a1a1a; line-height: 1.5; }}
        .header-band {{ width: 100%; background-color: #161616; margin-bottom: 16px; }}
        .header-band td {{ padding: 10px 14px; border: none; }}
        .school-name {{ color: #ffffff; font-size: 15pt; font-weight: bold; }}
        .app-name {{ color: #F8C300; font-size: 9pt; margin-top: 2px; }}
        h1 {{ font-size: 16pt; color: #CE1126; border-bottom: 2px solid #CE1126; padding-bottom: 4px; margin: 0 0 12px 0; }}
        h2 {{ font-size: 12.5pt; color: #161616; margin-top: 16px; }}
        h3 {{ font-size: 11pt; color: #161616; margin-top: 12px; }}
        p, li {{ font-size: 10.5pt; line-height: 1.5; }}
        ul, ol {{ margin: 4px 0 10px 0; }}
        hr {{ border: none; border-top: 1px solid #cccccc; margin: 18px 0 10px 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        th {{ background-color: #161616; color: white; padding: 6px 8px; text-align: left; font-size: 9.5pt; }}
        td {{ border: 1px solid #cccccc; padding: 6px 8px; font-size: 9.5pt; }}
        #footer_content {{ font-size: 8pt; color: #888888; text-align: center; }}
    </style>
    </head>
    <body>
    {header_band}
    <h1>{title}</h1>
    {body_html}
    <div id="footer_content">{school_name or "PNG Classroom Resource Generator"} - Page <pdf:pagenumber /></div>
    </body>
    </html>
    """
    buffer = io.BytesIO()
    pisa.CreatePDF(html, dest=buffer)
    return buffer.getvalue()


def markdown_to_docx_bytes(markdown_text: str, title: str, school_name: str | None = None) -> bytes:
    """Converts generated Markdown to a .docx a teacher can open and edit
    directly in Word, rather than just a flat, unstyled text file - headings,
    bold text, bullet lists, and tables all carry over as real Word
    formatting (via Markdown -> HTML -> docx), not just plain text.
    """
    text = markdown_text.replace("**Source(s):**", "---\n\n**Source(s):**")
    body_html = md.markdown(text, extensions=["tables", "nl2br"])

    document = Document()
    document.add_heading(title, level=1)
    if school_name:
        document.add_paragraph(school_name).runs[0].bold = True

    HtmlToDocx().add_html_to_document(body_html, document)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
